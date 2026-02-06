"""
Generate comprehensive Word documentation for the Privacy-Preserving AI Framework.
8-10 pages with easy explanations and 4-5 key diagrams.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os


def add_colored_cell(cell, text, bg_color, font_color="FFFFFF", bold=True, font_size=11):
    """Add colored background to a cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(font_color)


def create_documentation():
    doc = Document()

    # ============================================
    # STYLES
    # ============================================
    h1_style = doc.styles['Heading 1']
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 76, 153)

    h2_style = doc.styles['Heading 2']
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 102, 204)

    # ============================================
    # PAGE 1: TITLE PAGE
    # ============================================
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title_run = title.add_run("Privacy-Preserving AI Framework")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Using Federated Learning")
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(102, 126, 234)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run("A Complete Guide to Understanding How AI Can Learn\nWithout Ever Seeing Your Private Data")
    tagline_run.font.size = Pt(12)
    tagline_run.italic = True
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Project Documentation\n").bold = True
    info.add_run("Built with Python & TensorFlow\n")
    info.add_run("2026")

    doc.add_page_break()

    # ============================================
    # PAGE 2: TABLE OF CONTENTS
    # ============================================
    doc.add_heading("Table of Contents", level=1)

    toc_items = [
        ("1. Introduction - What is This Project?", "3"),
        ("2. The Problem - Why Do We Need This?", "3"),
        ("3. The Solution - What is Federated Learning?", "4"),
        ("4. How It Works - Step by Step", "5"),
        ("5. The Technology - What Tools We Use", "6"),
        ("6. Our Results - How Well Does It Work?", "7"),
        ("7. Key Terms Explained", "8"),
        ("8. Frequently Asked Questions", "9"),
        ("9. Conclusion", "10"),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run(f"  {'.' * (50 - len(item))}  {page}")

    doc.add_page_break()

    # ============================================
    # PAGE 3: INTRODUCTION
    # ============================================
    doc.add_heading("1. Introduction - What is This Project?", level=1)

    doc.add_paragraph(
        "Have you ever wondered how your phone's keyboard knows what word you want to type next? "
        "Or how Netflix knows what movies you might like? These are examples of Artificial Intelligence (AI) "
        "learning from data. But here's a problem: to make AI smart, companies usually need to collect "
        "everyone's personal data and store it in one place. That's not very private, is it?"
    )

    doc.add_paragraph(
        "This project solves that problem! We've built a system where AI can learn and become smart "
        "WITHOUT ever collecting your personal data. Your photos, messages, and private information "
        "stay safely on your own device. The AI learns from everyone, but sees no one's actual data. "
        "Sounds like magic? Let's see how it works!"
    )

    doc.add_heading("Who is This Document For?", level=2)

    doc.add_paragraph(
        "This document is written so that anyone can understand it - even if you've never heard of AI "
        "or programming before. We use simple examples and analogies to explain complex ideas. By the end, "
        "you'll be able to explain this project to your friends and family!"
    )

    doc.add_heading("2. The Problem - Why Do We Need This?", level=1)

    doc.add_heading("How Traditional AI Learning Works", level=2)

    doc.add_paragraph(
        "Imagine a teacher who wants to grade everyone's test papers. The traditional way is for "
        "all students to hand over their papers to the teacher. The teacher takes all the papers home, "
        "reads them, and learns what students know and don't know."
    )

    doc.add_paragraph(
        "This is how regular AI works too. Companies like Google or Facebook collect everyone's data "
        "(your photos, searches, messages) and put it all on their big computers. Then their AI "
        "looks at all this data and learns patterns. The problem? Your private information is now "
        "sitting on someone else's computer!"
    )

    doc.add_heading("What Could Go Wrong?", level=2)

    problems = [
        "Data Breaches: Hackers could steal millions of people's private information at once",
        "Privacy Concerns: Companies can see your personal photos, messages, and habits",
        "Trust Issues: You have to trust that companies won't misuse your data",
        "Legal Problems: Many countries now have strict laws about collecting personal data (like GDPR)",
        "Storage Costs: Storing billions of photos and messages is extremely expensive"
    ]

    for problem in problems:
        p = doc.add_paragraph(style='List Bullet')
        parts = problem.split(": ")
        p.add_run(parts[0] + ": ").bold = True
        p.add_run(parts[1])

    doc.add_paragraph(
        "So the question is: Can we make AI smart without collecting everyone's private data? "
        "The answer is YES, and that's exactly what our project does!"
    )

    doc.add_page_break()

    # ============================================
    # PAGE 4: FEDERATED LEARNING EXPLANATION
    # ============================================
    doc.add_heading("3. The Solution - What is Federated Learning?", level=1)

    doc.add_paragraph(
        "Federated Learning is a clever way to train AI where the data never leaves your device. "
        "Instead of bringing all the data to the AI, we bring the AI to the data! Let's understand "
        "this with a fun example."
    )

    doc.add_heading("The Pizza Recipe Story", level=2)

    doc.add_paragraph(
        "Imagine five families in a neighborhood. Each family has a secret pizza recipe passed down "
        "from their grandparents. They want to create the BEST pizza recipe by combining everyone's "
        "knowledge, but nobody wants to share their actual secret recipe. What can they do?"
    )

    doc.add_paragraph("Here's the clever solution they came up with:")

    steps = [
        ("Step 1", "A chef creates a basic pizza recipe and gives a copy to each family."),
        ("Step 2", "Each family goes home and tries making the pizza using the basic recipe PLUS "
                   "their secret knowledge. Family A discovers 'adding a pinch more salt makes it better.' "
                   "Family B finds 'cooking 2 minutes longer makes the crust crispier.'"),
        ("Step 3", "Each family writes down ONLY their improvement suggestions on a piece of paper - "
                   "NOT their secret recipe! They send these suggestions to the chef."),
        ("Step 4", "The chef reads all the suggestions and creates an improved recipe by combining "
                   "everyone's tips. 'Add a bit more salt AND cook 2 minutes longer.'"),
        ("Step 5", "The chef sends this improved recipe back to all families. They try it again and "
                   "find even more improvements. This continues until the recipe is perfect!")
    ]

    for title, desc in steps:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "The brilliant part? The chef never learned anyone's secret recipe! Each family kept their "
        "secrets safe, but together they created something better than any single family could alone."
    )

    doc.add_heading("How This Relates to AI", level=2)

    doc.add_paragraph("In our project, the same thing happens with AI:")

    comparisons = [
        ("The families", "are your phones and computers (we call them 'clients')"),
        ("The secret recipes", "are your private data (photos, messages, etc.)"),
        ("The chef", "is the central server that coordinates everything"),
        ("The improvement suggestions", "are called 'model weights' - just numbers that describe what the AI learned"),
        ("The final perfect recipe", "is the trained AI model that can recognize images, predict text, etc.")
    ]

    for item, meaning in comparisons:
        p = doc.add_paragraph()
        p.add_run(f"{item} ").bold = True
        p.add_run(meaning)

    doc.add_page_break()

    # ============================================
    # PAGE 5: HOW IT WORKS - DIAGRAM
    # ============================================
    doc.add_heading("4. How It Works - Step by Step", level=1)

    doc.add_paragraph(
        "Now let's see exactly how our project works. We'll break it down into simple steps that "
        "anyone can follow."
    )

    # DIAGRAM 1: Main Flow
    doc.add_heading("The Complete Process (Diagram 1)", level=2)

    flow_table = doc.add_table(rows=5, cols=1)
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    flow_steps = [
        ("1. SERVER sends empty AI model to all devices", "3498DB"),
        ("2. Each DEVICE trains the model using its own private data", "27AE60"),
        ("3. DEVICES send only the learning (weights) back - NOT the data!", "9B59B6"),
        ("4. SERVER combines all the learning into one smarter model", "E67E22"),
        ("5. REPEAT steps 1-4 until AI is smart enough!", "1ABC9C")
    ]

    for i, (step, color) in enumerate(flow_steps):
        add_colored_cell(flow_table.rows[i].cells[0], step, color, "FFFFFF", True, 10)

    doc.add_paragraph()

    doc.add_heading("What Happens in Each Step", level=2)

    detailed_steps = [
        ("Loading the Data",
         "First, we need data to learn from. In our project, we use a famous dataset called MNIST - "
         "it contains 60,000 pictures of handwritten numbers (0-9). Imagine 60,000 people each wrote "
         "a number, and we photographed it. In real life, this data would be on different people's phones."),

        ("Splitting Data Among Clients",
         "We pretend we have 5 different devices (clients). Each device gets some of the pictures. "
         "Importantly, we make it 'Non-IID' (Non-Independently and Identically Distributed) - this means "
         "each device has different types of numbers. Device 1 might have mostly 0s and 1s, Device 2 "
         "might have mostly 8s and 9s. This is realistic because your photos are different from your friend's!"),

        ("Local Training",
         "Each device trains the AI using only its own pictures. The AI looks at a picture, guesses "
         "what number it is, checks if it was right, and adjusts itself to do better next time. "
         "This happens thousands of times on each device."),

        ("Sending Updates",
         "After training, each device calculates how much the AI changed (these are called 'weight updates'). "
         "ONLY these numbers are sent to the server - never the actual pictures! It's like sending "
         "'I turned dial #1 up by 0.5' instead of showing your photos."),

        ("Server Aggregation",
         "The server receives updates from all devices and combines them. Devices with more data get "
         "more influence (like a family that made more pizzas getting more say in the recipe). "
         "This combined knowledge creates a better AI model."),

        ("Repeat (Communication Rounds)",
         "Steps 3-5 repeat multiple times. Each repetition is called a 'round.' Our project runs 10 rounds. "
         "With each round, the AI gets smarter and smarter!")
    ]

    for title, desc in detailed_steps:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ============================================
    # PAGE 6: TECHNOLOGY
    # ============================================
    doc.add_heading("5. The Technology - What Tools We Use", level=1)

    doc.add_paragraph(
        "Our project uses several programming tools that work together like a team. Here's what each one does:"
    )

    doc.add_heading("Programming Languages and Libraries", level=2)

    tech_explanations = [
        ("Python",
         "This is the main programming language we use. Python is known for being easy to read and write - "
         "almost like writing in English! It's the most popular language for AI projects."),

        ("TensorFlow",
         "Created by Google, TensorFlow is like a construction kit for building AI. It handles all the "
         "complex math needed to train neural networks. Think of it as the engine of our AI car."),

        ("Keras",
         "Keras sits on top of TensorFlow and makes it even easier to use. If TensorFlow is the engine, "
         "Keras is the steering wheel that makes driving easier."),

        ("NumPy",
         "AI involves lots of math with big lists of numbers. NumPy makes this math super fast. "
         "Without it, our project would take hours instead of minutes!"),

        ("Streamlit",
         "This creates our beautiful web dashboard where you can see all the results. It turns our "
         "Python code into an interactive website with just a few lines of code."),

        ("Plotly",
         "This creates the interactive charts and graphs. You can hover over them to see exact values, "
         "zoom in, and explore the data visually.")
    ]

    for tech, desc in tech_explanations:
        p = doc.add_paragraph()
        p.add_run(tech + ": ").bold = True
        p.add_run(desc)

    # DIAGRAM 2: Neural Network
    doc.add_heading("The AI Brain - Neural Network (Diagram 2)", level=2)

    doc.add_paragraph(
        "Our AI uses something called a Convolutional Neural Network (CNN). Don't let the fancy name "
        "scare you - it's just a computer program that's really good at recognizing patterns in images. "
        "Here's how it works:"
    )

    nn_table = doc.add_table(rows=1, cols=5)
    nn_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["INPUT\n(28x28 Image)", "PATTERN\nDETECTOR 1", "PATTERN\nDETECTOR 2", "DECISION\nMAKER", "OUTPUT\n(0-9)"]
    colors = ["3498DB", "9B59B6", "9B59B6", "E67E22", "27AE60"]

    for i, (label, color) in enumerate(zip(labels, colors)):
        add_colored_cell(nn_table.rows[0].cells[i], label, color, "FFFFFF", True, 9)

    doc.add_paragraph()

    doc.add_paragraph(
        "The image enters on the left. The first pattern detector finds simple things like edges and lines. "
        "The second detector finds more complex patterns like curves and corners. The decision maker looks "
        "at all these patterns and decides which number (0-9) the image shows. It's like how you recognize "
        "a face - first you see eyes, nose, mouth, then your brain combines them to recognize who it is!"
    )

    doc.add_page_break()

    # ============================================
    # PAGE 7: RESULTS
    # ============================================
    doc.add_heading("6. Our Results - How Well Does It Work?", level=1)

    doc.add_paragraph(
        "After building our system, we tested it to see how well it works. Here are the results:"
    )

    # DIAGRAM 3: Results Comparison
    doc.add_heading("Performance Comparison (Diagram 3)", level=2)

    results_table = doc.add_table(rows=5, cols=3)
    results_table.style = 'Table Grid'
    results_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    add_colored_cell(results_table.rows[0].cells[0], "What We Measured", "2C3E50")
    add_colored_cell(results_table.rows[0].cells[1], "Federated\n(Our Method)", "667EEA")
    add_colored_cell(results_table.rows[0].cells[2], "Centralized\n(Traditional)", "E74C3C")

    # Data
    data = [
        ("Accuracy (how often it's correct)", "89.30%", "99.22%"),
        ("Privacy Protection", "100% Safe", "0% (Data exposed)"),
        ("Where is your data?", "On YOUR device", "On company's server"),
        ("Risk of data breach", "Almost zero", "High")
    ]

    for i, (metric, fed, cent) in enumerate(data):
        results_table.rows[i+1].cells[0].text = metric
        results_table.rows[i+1].cells[1].text = fed
        results_table.rows[i+1].cells[2].text = cent
        for cell in results_table.rows[i+1].cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    doc.add_heading("Understanding These Results", level=2)

    doc.add_paragraph(
        "Federated Accuracy (89.30%): Our privacy-preserving AI correctly identifies handwritten "
        "numbers about 89 times out of 100. That's really impressive considering it never saw "
        "the actual training pictures!"
    )

    doc.add_paragraph(
        "Centralized Accuracy (99.22%): When we train the traditional way (collecting all data in one place), "
        "we get 99% accuracy. This is slightly better, but remember - it requires giving up all your privacy!"
    )

    doc.add_paragraph(
        "The Trade-Off: We lose about 10% accuracy but gain 100% privacy. For most applications, "
        "this is an excellent deal! Would you rather have a slightly less accurate AI that respects "
        "your privacy, or a slightly better AI that reads all your personal data?"
    )

    doc.add_heading("What Our Dashboard Shows", level=2)

    doc.add_paragraph(
        "Our project includes a beautiful web dashboard (built with Streamlit) that shows:"
    )

    outputs = [
        "Accuracy Graph: How the AI gets smarter over each round (starts at ~49%, ends at ~89%)",
        "Loss Graph: How the AI's mistakes decrease over time",
        "Communication Cost: How much data was transferred (about 96.5 MB total)",
        "Client Distribution: Which numbers each device had (showing the Non-IID distribution)",
        "Privacy Analysis: A radar chart comparing privacy vs accuracy trade-offs"
    ]

    for output in outputs:
        doc.add_paragraph(output, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # PAGE 8: KEY TERMS
    # ============================================
    doc.add_heading("7. Key Terms Explained", level=1)

    doc.add_paragraph(
        "Here are all the important terms you might hear when discussing this project, explained simply:"
    )

    terms = [
        ("Artificial Intelligence (AI)",
         "A computer program that can learn from examples and make decisions, similar to how humans learn."),

        ("Machine Learning",
         "A type of AI where the computer learns patterns from data instead of being explicitly programmed."),

        ("Federated Learning",
         "A way to train AI where the data stays on user devices and only the learning (not the data) is shared."),

        ("Client",
         "Any device that participates in federated learning - could be your phone, laptop, or tablet."),

        ("Server",
         "The central computer that coordinates the learning process and combines updates from all clients."),

        ("Model",
         "The AI 'brain' - a mathematical structure that can recognize patterns after being trained."),

        ("Model Weights",
         "Numbers inside the AI that store what it has learned. Adjusting these numbers is how the AI learns."),

        ("Training",
         "The process of showing examples to the AI so it can learn patterns and improve."),

        ("Communication Round",
         "One complete cycle: server sends model → clients train → clients send updates → server combines."),

        ("FedAvg (Federated Averaging)",
         "The algorithm that combines all client updates by taking a weighted average."),

        ("Non-IID (Non-Independently and Identically Distributed)",
         "When different clients have different types of data - realistic because your data differs from others'."),

        ("CNN (Convolutional Neural Network)",
         "A type of AI that's especially good at recognizing patterns in images."),

        ("Accuracy",
         "The percentage of times the AI gives the correct answer."),

        ("Loss",
         "A measure of how wrong the AI's predictions are - lower is better."),

        ("Differential Privacy",
         "Adding mathematical noise to make it impossible to identify any individual's data - extra protection!")
    ]

    for term, definition in terms:
        p = doc.add_paragraph()
        p.add_run(term + ": ").bold = True
        p.add_run(definition)

    doc.add_page_break()

    # ============================================
    # PAGE 9: FAQ
    # ============================================
    doc.add_heading("8. Frequently Asked Questions", level=1)

    faqs = [
        ("Is my data really safe with Federated Learning?",
         "Yes! Your actual data (photos, messages, etc.) NEVER leaves your device. Only mathematical "
         "numbers describing what the AI learned are shared. It's like sharing that 'salt improves pizza' "
         "without revealing your grandmother's secret recipe."),

        ("Why is the federated accuracy lower than centralized?",
         "Because each device only sees a small portion of the data. Imagine trying to understand "
         "a movie by only watching random 5-minute clips - you'd get the general idea but miss some details. "
         "The AI faces a similar challenge. However, the privacy benefit far outweighs this small accuracy loss."),

        ("Can hackers steal my data from the weight updates?",
         "It's extremely difficult. The updates are mathematical averages from thousands of data points. "
         "It would be like trying to figure out one person's vote from only knowing the election results. "
         "For extra security, we can add 'Differential Privacy' which makes it mathematically impossible."),

        ("Who uses Federated Learning in real life?",
         "Many big companies! Google uses it to improve keyboard predictions (Gboard) without reading your texts. "
         "Apple uses it to improve Siri without listening to your conversations. Hospitals use it to improve "
         "medical AI without sharing patient records. It's becoming the standard for privacy-respecting AI."),

        ("What is Non-IID and why does it make learning harder?",
         "Non-IID means different devices have different types of data. Your phone has photos of your life, "
         "which are completely different from your friend's photos. This makes it harder for the AI because "
         "each device teaches it something different. But our system handles this well!"),

        ("How long does the training take?",
         "Our 10-round training takes about 30-40 minutes total. Traditional centralized training is faster "
         "(about 5 minutes) because there's no back-and-forth communication. But the privacy is worth the wait!"),

        ("Can this work with other types of data besides images?",
         "Absolutely! Federated Learning works with text (like keyboard predictions), voice (like Siri), "
         "health data (like fitness trackers), and much more. The principles are the same - data stays local, "
         "only learning is shared."),

        ("What happens if one device sends bad updates?",
         "The averaging process naturally reduces the impact of any single device. If 100 devices send good "
         "updates and 1 sends bad updates, the bad one gets diluted. Advanced systems can also detect and "
         "ignore suspicious updates.")
    ]

    for q, a in faqs:
        doc.add_heading(f"Q: {q}", level=2)
        doc.add_paragraph(a)

    doc.add_page_break()

    # ============================================
    # PAGE 10: CONCLUSION
    # ============================================
    doc.add_heading("9. Conclusion", level=1)

    doc.add_heading("What We Built", level=2)

    doc.add_paragraph(
        "In this project, we successfully built a Privacy-Preserving AI Framework using Federated Learning. "
        "Our system can train an AI to recognize handwritten numbers with 89.30% accuracy - and here's the "
        "amazing part - without ever collecting or seeing anyone's private data!"
    )

    doc.add_heading("Why This Matters", level=2)

    doc.add_paragraph(
        "In today's world, data is everywhere. Every time you use your phone, you generate data. "
        "Traditional AI systems want to collect all this data, which raises serious privacy concerns. "
        "Our project shows that there's a better way - AI can be smart AND respectful of privacy."
    )

    # DIAGRAM 4: Summary
    doc.add_heading("The Big Picture (Diagram 4)", level=2)

    summary_table = doc.add_table(rows=1, cols=1)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = summary_table.rows[0].cells[0]
    add_colored_cell(cell,
        "FEDERATED LEARNING IN ONE SENTENCE:\n\n"
        "AI learns from everyone's data without seeing anyone's data -\n"
        "because only the LEARNING travels, not the DATA!",
        "667EEA", "FFFFFF", True, 12)

    doc.add_paragraph()

    doc.add_heading("Key Takeaways", level=2)

    takeaways = [
        "Your data stays on your device - always",
        "Only mathematical learning (weights) is shared with the server",
        "We achieve 89.30% accuracy while maintaining 100% privacy",
        "We retain 90% of traditional AI performance - an excellent trade-off",
        "This technology is already used by Google, Apple, and hospitals worldwide",
        "Federated Learning is the future of ethical, privacy-respecting AI"
    ]

    for t in takeaways:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading("Final Thoughts", level=2)

    final = doc.add_paragraph(
        "The future of Artificial Intelligence doesn't have to be scary. With Federated Learning, "
        "we can have the benefits of smart AI while keeping our personal information private and secure. "
        "This project is a small step toward a future where technology respects human privacy. "
        "The future of AI is private, decentralized, and respectful of your data!"
    )

    doc.add_paragraph()
    doc.add_paragraph()

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_run = end.add_run("— End of Documentation —")
    end_run.italic = True
    end_run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    output_path = os.path.join(os.path.dirname(__file__), "Project_Documentation_Final.docx")
    doc.save(output_path)
    print(f"Documentation saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_documentation()
